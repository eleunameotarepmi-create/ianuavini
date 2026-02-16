
import docx
import sys

def extract_text(doc_path, out_path):
    try:
        doc = docx.Document(doc_path)
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())
        
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    full_text.append(" | ".join(row_text))

        with open(out_path, 'w', encoding='utf-8') as f:
            f.write("\n".join(full_text))
            
        print(f"Successfully wrote to {out_path}")
    except Exception as e:
        print(f"Error: {e}")

if __name__ == "__main__":
    path = r"C:\Surface Shares\OneDrive - MSFT\Documenti\LAVORO\Schede vini\CantinePiemonte\Zone Vinicole e sottozone piemonte.docx"
    out = "piemonte_zones.txt"
    extract_text(path, out)
